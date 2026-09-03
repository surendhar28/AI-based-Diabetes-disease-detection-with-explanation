import re
import base64
import urllib.request
import io
import time
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def parse_inline_styles(paragraph, text):
    # Regex to find ***bold-italic***, **bold**, *italic*, and `code` patterns
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # soft red-crimson for inline code
        else:
            # Clean up raw mathematical brackets if any, or preserve standard text
            paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set premium Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base typography styling
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 37, 41) # Dark charcoal instead of pure black for readability
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    is_mermaid_block = False
    code_block_text = []
    last_heading_text = "Architecture Diagram"
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code block boundaries
        if stripped.startswith('```'):
            if in_code_block:
                # Close code block
                in_code_block = False
                if is_mermaid_block:
                    mermaid_code = '\n'.join(code_block_text)
                    try:
                        max_retries = 3
                        image_data = None
                        for attempt in range(max_retries):
                            try:
                                print(f"Attempting to render mermaid diagram for: {last_heading_text} (Attempt {attempt+1}/{max_retries})...")
                                graph_bytes = mermaid_code.encode("utf-8")
                                base64_string = base64.urlsafe_b64encode(graph_bytes).decode("ascii").strip("=")
                                url = f"https://mermaid.ink/img/{base64_string}"
                                
                                req = urllib.request.Request(
                                    url, 
                                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                                )
                                with urllib.request.urlopen(req, timeout=25) as response:
                                    image_data = response.read()
                                break
                            except Exception as e:
                                if attempt == max_retries - 1:
                                    raise e
                                print(f"Attempt {attempt+1} failed ({e}). Retrying in 2 seconds...")
                                time.sleep(2)
                        
                        image_stream = io.BytesIO(image_data)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(4)
                        run = p.add_run()
                        run.add_picture(image_stream, width=Inches(5.5))
                        
                        # Add a caption
                        p_caption = doc.add_paragraph()
                        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_caption.paragraph_format.space_before = Pt(2)
                        p_caption.paragraph_format.space_after = Pt(12)
                        p_caption.paragraph_format.keep_with_next = True
                        run_cap = p_caption.add_run(f"Figure: {last_heading_text}")
                        run_cap.italic = True
                        run_cap.font.size = Pt(9.5)
                        run_cap.font.color.rgb = RGBColor(128, 128, 128)
                        print("Diagram successfully embedded.")
                    except Exception as exc:
                        print(f"Warning: Failed to render mermaid diagram ({exc}). Falling back to text code block.")
                        # Fallback to normal code block rendering
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.4)
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(6)
                        
                        run = p.add_run(mermaid_code)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9.0)
                        run.font.color.rgb = RGBColor(60, 60, 60)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(6)
                    
                    code_text = '\n'.join(code_block_text)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.0)
                    run.font.color.rgb = RGBColor(60, 60, 60)
                
                code_block_text = []
                is_mermaid_block = False
            else:
                in_code_block = True
                lang = stripped[3:].strip()
                if lang == 'mermaid':
                    is_mermaid_block = True
                else:
                    is_mermaid_block = False
            i += 1
            continue
            
        if in_code_block:
            code_block_text.append(line.rstrip('\n'))
            i += 1
            continue
            
        # Headers
        if stripped.startswith('# '):
            last_heading_text = stripped[2:].strip()
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            parse_inline_styles(p, stripped[2:])
        elif stripped.startswith('## '):
            last_heading_text = stripped[3:].strip()
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            parse_inline_styles(p, stripped[3:])
        elif stripped.startswith('### '):
            last_heading_text = stripped[4:].strip()
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            parse_inline_styles(p, stripped[4:])
        elif stripped.startswith('#### '):
            last_heading_text = stripped[5:].strip()
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            parse_inline_styles(p, stripped[5:])
        # Horizontal rules
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run('__________________________________________________________________')
            run.font.color.rgb = RGBColor(210, 214, 219)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Callouts / Alert blocks
        elif stripped.startswith('> '):
            first_line = stripped[2:]
            alert_type = ''
            if first_line.startswith('[!'):
                match = re.match(r'^\[\!(.*?)\]', first_line)
                if match:
                    alert_type = match.group(1)
                    first_line = first_line[len(match.group(0)):].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            if alert_type:
                run = p.add_run(f"[{alert_type.upper()}] ")
                run.bold = True
                if alert_type.upper() in ['IMPORTANT', 'WARNING', 'CAUTION']:
                    run.font.color.rgb = RGBColor(199, 37, 78) # high-alert red
                else:
                    run.font.color.rgb = RGBColor(0, 102, 204) # info blue
            
            quote_lines = [first_line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('> '):
                line_text = lines[i].strip()[2:]
                if not (line_text.startswith('[!') and ']' in line_text):
                    quote_lines.append(line_text)
                i += 1
            i -= 1
            full_quote = ' '.join(quote_lines).strip()
            parse_inline_styles(p, full_quote)
        # Images syntax ![alt](path)
        elif stripped.startswith('![') and '](' in stripped and stripped.endswith(')'):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
            if match:
                alt_text, img_rel_path = match.group(1), match.group(2)
                # Resolve image path relative to md file or workspace root
                md_dir = Path(md_path).parent
                img_path = (md_dir / img_rel_path).resolve()
                if not img_path.exists():
                    # Try workspace root fallback
                    img_path = (Path(__file__).resolve().parent / img_rel_path).resolve()
                
                if img_path.exists():
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(4)
                        run = p.add_run()
                        run.add_picture(str(img_path), width=Inches(5.8))
                        
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(2)
                        p_cap.paragraph_format.space_after = Pt(12)
                        run_cap = p_cap.add_run(f"Figure: {alt_text if alt_text else 'Illustration'}")
                        run_cap.italic = True
                        run_cap.font.size = Pt(9.5)
                        run_cap.font.color.rgb = RGBColor(100, 116, 139)
                        print(f"Successfully embedded image: {img_path.name}")
                    except Exception as e:
                        print(f"Warning: Failed to insert image {img_path}: {e}")
                else:
                    print(f"Warning: Image file not found at {img_path}")
        # Markdown Tables (| col1 | col2 |)
        elif stripped.startswith('|') and stripped.endswith('|'):
            table_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1
            
            # Parse table rows
            rows_data = []
            for t_line in table_lines:
                # Skip separator line like |---|---|
                if re.match(r'^\|[\s:-|-]+\|$', t_line):
                    continue
                cells = [c.strip() for c in t_line.strip('|').split('|')]
                rows_data.append(cells)
                
            if rows_data:
                num_rows = len(rows_data)
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                table.autofit = True
                
                for r_idx, row_cells in enumerate(rows_data):
                    row = table.rows[r_idx]
                    is_header = (r_idx == 0)
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            if is_header:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                # Dark slate background for header
                                shading_elm = parse_xml(r'<w:shd {} w:fill="1E293B"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading_elm)
                            else:
                                parse_inline_styles(p, cell_text)
                                if r_idx % 2 == 1:
                                    shading_elm = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                                    cell._tc.get_or_add_tcPr().append(shading_elm)
                # Add spacing after table
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(6)
        # Bullet list items
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_styles(p, stripped[2:])
        # Numbered list items
        elif re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^\d+\.\s', stripped)
            offset = len(match.group(0))
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_styles(p, stripped[offset:])
        # Blank lines
        elif not stripped:
            pass
        # Normal paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parse_inline_styles(p, stripped)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully converted MD to DOCX at {docx_path}")

if __name__ == '__main__':
    import sys
    md = sys.argv[1] if len(sys.argv) > 1 else 'EXPLANATION.md'
    docx = sys.argv[2] if len(sys.argv) > 2 else 'EXPLANATION.docx'
    convert_md_to_docx(md, docx)

